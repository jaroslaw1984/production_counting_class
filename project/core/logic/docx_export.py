import tempfile
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from project.config.paths import BASE_DIR

# --- funkcje eksportu/drukowania/edycji raportu DOCX ---       
def export_report_docx(report_data: dict, template_path: Path | None = None) -> Path:
    """Generuje DOCX na bazie template. Zwraca ścieżkę do wygenerowanego pliku."""
    if not report_data:
        raise ValueError("Brak danych raportu (report_data).")

    if template_path is None:
        template_path = BASE_DIR / "templates" / "report_template.docx"

    if not template_path.exists():
        raise FileNotFoundError(f"Brak szablonu DOCX: {template_path}")
    
    doc = Document(str(template_path))
    
    def set_cell_margins(cell, top=200, bottom=200):
        """
        Marginesy w TWIPS:
        200 ≈ ok. 0.35 cm
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        tcMar = OxmlElement("w:tcMar")

        for side, value in [("top", top), ("bottom", bottom)]:
            node = OxmlElement(f"w:{side}")
            node.set(qn("w:w"), str(value))
            node.set(qn("w:type"), "dxa")
            tcMar.append(node)

        tcPr.append(tcMar)        

    def replace_all(old: str, new: str) -> None:
        # paragrafy
        for p in doc.paragraphs:
            if old in p.text:
                for r in p.runs:
                    r.text = r.text.replace(old, new)
        # tabele
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    if old in cell.text:
                        cell.text = cell.text.replace(old, new)

    # 1) Podmień pola nagłówka
    replace_all("{{SHIFT_INFO}}", str(report_data.get("shift_info", "")))
    replace_all("{{REPORT_DATE}}", str(report_data.get("report_date", "")))
    replace_all("{{USER}}", str(report_data.get("user", "")))
    replace_all("{{LINE}}", str(report_data.get("line", "")))
    replace_all("{{MACHINE}}", str(report_data.get("machine", "")))
    replace_all("{{PALLETS_TOTAL}}", str(report_data.get("pallets_total", "")))

    # 2) Wypełnij tabelę: znajdź wiersz-placeholder i zastąp go danymi
    rows = report_data.get("rows", []) or []
    if not rows:
        # nic do tabeli, zostaw placeholdery jako puste
        replace_all("{{ROW_LP}}", "")
        replace_all("{{ROW_INDEX}}", "")
        replace_all("{{ROW_QTY_M}}", "")
        replace_all("{{ROW_PCS}}", "")
        replace_all("{{ROW_PALLETS}}", "")
    else:
        # szukamy pierwszej tabeli z placeholderem
        target_table = None
        placeholder_row_idx = None
        for t in doc.tables:
            for i, r in enumerate(t.rows):
                if any("{{ROW_LP}}" in c.text for c in r.cells):
                    target_table = t
                    placeholder_row_idx = i
                    break
            if target_table is not None:
                break

        if target_table is None or placeholder_row_idx is None:
            raise ValueError("Nie znalazłem w szablonie wiersza placeholderów ({{ROW_LP}}...).")

        # usuń wiersz placeholder
        tbl = target_table._tbl
        tr = target_table.rows[placeholder_row_idx]._tr
        tbl.remove(tr)

        # dodaj wiersze danych
        for item in rows:
            r = target_table.add_row().cells
            r[0].text = str(item.get("lp", ""))
            r[1].text = str(item.get("index", ""))
            r[2].text = str(item.get("qty_m", ""))
            r[3].text = str(item.get("pcs", ""))
            r[4].text = str(item.get("pallets", ""))
            
            for cell in r:
                set_cell_margins(cell, top=200, bottom=200)

    # 3) Zapis
    out_dir = Path(tempfile.gettempdir()) / "production_counter_reports"
    out_dir.mkdir(parents=True, exist_ok=True)
    stamp = datetime.now().strftime("%Y-%m-%d_%H%M%S")
    safe_line = str(report_data.get("line", "LINIA")).replace("/", "-")
    out_path = out_dir / f"Zapotrzebowanie_{safe_line}_{stamp}.docx"
    doc.save(str(out_path))
    return out_path